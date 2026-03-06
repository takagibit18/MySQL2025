import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_font_dengxian(run):
    """辅助函数：将Run的字体设置为等线"""
    run.font.name = '等线'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')


def add_heading_with_color(doc, text, level, color=None):
    """添加带颜色的标题，并强制字体为等线"""
    heading = doc.add_heading(text, level)
    # 遍历标题中的所有run，设置字体和颜色
    for run in heading.runs:
        set_font_dengxian(run)  # 设置字体为等线
        if color:
            run.font.color.rgb = color
    return heading


def add_text_with_color(doc, text, color=None, bold=False):
    """添加带颜色的文本（正文默认宋体）"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    if color:
        run.font.color.rgb = color
    if bold:
        run.bold = True
    return para


def add_image_placeholder(doc, description, color=RGBColor(255, 0, 0)):
    """添加图片占位符"""
    para = doc.add_paragraph()
    run = para.add_run(f"[图片占位符：{description}]")
    run.font.color.rgb = color
    run.font.size = Pt(12)
    run.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para


def create_report():
    """生成实验报告第三部分"""
    doc = Document()

    # --- 全局设置：正文默认为宋体 ---
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ================= 标题 (Level 0) =================
    title = doc.add_heading('第三部分  学生管理系统开发', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 强制主标题字体为等线
    for run in title.runs:
        set_font_dengxian(run)
        run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色

    # ================= 一、需求分析 =================
    add_heading_with_color(doc, '一、需求分析', 1)
    doc.add_paragraph(
        '本系统是一个基于Django框架开发的综合性学生管理平台，设计参考了机票预订系统的分层架构。'
        '系统旨在为学校提供全面的数字化管理解决方案，涵盖学生、教师、管理员三个核心角色的业务流程。'
    )
    doc.add_paragraph('核心需求包括：')
    doc.add_paragraph('（1）多角色权限管理：实现学生、教师、管理员的独立登录与功能隔离，基于Session进行会话管理。',
                      style='List Bullet')
    doc.add_paragraph('（2）基础教务管理：支持课程管理、班级管理、选课退课、成绩录入（含Excel导入）及查询。',
                      style='List Bullet')
    doc.add_paragraph('（3）考勤管理模块（自主设计）：这是系统的核心特色功能。需求要求实现从“教师课堂点名”到“学生请假申请”再到“教师审批”的完整闭环。'
                      '特别需要支持病假条图片的上传与预览、考勤状态的批量操作以及审批通过后的自动考勤记录生成。')

    # ================= 二、系统功能设计 =================
    add_heading_with_color(doc, '二、系统功能设计', 1)

    # 2.1 系统功能结构
    add_heading_with_color(doc, '1. 系统功能结构', 2)
    doc.add_paragraph('系统采用分离式多应用设计，主要功能模块如下：')
    add_image_placeholder(doc, '系统功能结构图（展示Student Portal, Teacher Portal, Admin Portal及其子功能）',
                          RGBColor(255, 0, 0))

    doc.add_paragraph('各门户主要功能：')
    doc.add_paragraph('• 学生门户：个人信息、选课管理、成绩查询、考勤查询、请假申请（含图片上传）。')
    doc.add_paragraph('• 教师门户：课程管理、查看学生名单、成绩录入、考勤点名（交互式）、请假审批。')
    doc.add_paragraph('• 管理员门户：学生/教师/班级/课程的增删改查、认证授权管理。')

    # 2.2 系统业务流程
    add_heading_with_color(doc, '2. 系统业务流程', 2)
    doc.add_paragraph('以核心的“考勤与请假”业务为例，流程设计如下：')
    add_image_placeholder(doc, '考勤业务流程图（学生提交申请 -> 教师审批 -> 自动写入考勤表 -> 更新状态）',
                          RGBColor(0, 0, 255))

    doc.add_paragraph('关键流程逻辑：')
    doc.add_paragraph('1. 点名流程：教师选择课程/日期 -> 加载学生 -> 点击头像切换状态/批量全选 -> 保存 -> 数据库更新。',
                      style='List Bullet')
    doc.add_paragraph('2. 请假流程：学生填写表单(含文件) -> 存入leave_requests -> 状态“待审批”。', style='List Bullet')
    doc.add_paragraph(
        '3. 审批流程：教师查看详情(含图片) -> 通过/拒绝 -> 若通过，系统自动计算日期范围，在attendances表中创建对应日期的“请假”记录。',
        style='List Bullet')

    # ================= 三、系统开发 =================
    add_heading_with_color(doc, '三、系统开发', 1)

    # 3.1 系统开发环境
    add_heading_with_color(doc, '1. 系统开发环境', 2)
    env_items = [
        '开发语言：Python 3.8+',
        'Web框架：Django 5.0.7',
        '数据库：MySQL 8.0 (使用mysqlclient驱动，原生SQL操作)',
        '前端框架：Bootstrap 5.3.3',
        '交互技术：原生JavaScript + Fetch API (AJAX异步请求)',
        '文件存储：Django Media配置 (用于存储病假条等)'
    ]
    for item in env_items:
        doc.add_paragraph(f'（{env_items.index(item) + 1}）{item}', style='List Bullet')

    # 3.2 文件夹组织结构
    add_heading_with_color(doc, '2. 文件夹组织结构', 2)
    doc.add_paragraph('项目遵循Django多应用架构，文件结构如下：')

    # 模拟代码块格式
    tree_text = """student_management_system/
├── student_site/             # 主项目配置 (settings.py, urls.py)
├── student_portal/           # 学生端应用 (views.py: 登录, 选课, 考勤)
├── teacher_portal/           # 教师端应用 (views.py: 成绩, 点名, 审批)
├── admin_portal/             # 管理员端应用 (views.py, forms.py)
├── templates/                # HTML模板文件
│   ├── base.html             # 基础模板
│   ├── student/              # 学生端模板 (attendance.html, grades.html...)
│   ├── teacher/              # 教师端模板 (attendance.html, leave_requests.html...)
│   └── admin/                # 管理员端模板
├── static/                   # 静态资源 (CSS, JS, img)
├── media/                    # 媒体文件 (medical_certificates/)
├── database.sql              # 基础建表SQL
└── 考勤模块数据库扩展.sql     # 考勤模块专用SQL"""

    p = doc.add_paragraph(tree_text)
    p.style.font.name = 'Courier New'
    p.style.font.size = Pt(9)

    add_image_placeholder(doc, 'PyCharm项目目录结构截图', RGBColor(0, 128, 0))

    # ================= 四、数据库设计 =================
    add_heading_with_color(doc, '四、数据库设计', 1)

    # 4.1 数据表模型
    add_heading_with_color(doc, '1. 数据表模型', 2)
    doc.add_paragraph(
        '系统包含8个核心表和1个扩展表。核心表包括students, teachers, admins, classes, courses, enrollments, grades。')
    doc.add_paragraph('针对考勤模块，特别设计了以下数据表结构：')

    # 考勤表
    # 注意：此处为Heading 3级别
    h3_1 = add_heading_with_color(doc, '（1）考勤表 (attendances) - 包含扩展字段', 3)

    table1 = doc.add_table(rows=1, cols=4)
    table1.style = 'Light Grid Accent 1'
    hdr1 = table1.rows[0].cells
    hdr1[0].text = '字段名'
    hdr1[1].text = '类型'
    hdr1[2].text = '约束'
    hdr1[3].text = '说明'

    att_data = [
        ('考勤ID', 'CHAR(15)', 'PK', '主键'),
        ('学号', 'CHAR(12)', 'FK', '关联students'),
        ('课程ID', 'CHAR(10)', 'FK', '关联courses'),
        ('考勤日期', 'DATE', 'NOT NULL', '记录日期'),
        ('考勤状态', 'TINYINT', 'NOT NULL', '1-出勤, 2-迟到, 3-早退, 4-缺勤, 5-请假'),
        ('备注', 'VARCHAR(200)', '', '行内编辑备注'),
        ('记录教师工号', 'CHAR(10)', 'FK', '关联teachers'),
        ('请假申请ID', 'CHAR(15)', 'FK', '关联leave_requests (扩展)'),
        ('审批状态', 'TINYINT', '', '0-未审批, 1-通过, 2-拒绝 (冗余)'),
    ]
    for row in att_data:
        cells = table1.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val

    doc.add_paragraph('')

    # 请假申请表
    h3_2 = add_heading_with_color(doc, '（2）请假申请表 (leave_requests) - 考勤模块专用', 3)

    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Light Grid Accent 1'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = '字段名'
    hdr2[1].text = '类型'
    hdr2[2].text = '约束'
    hdr2[3].text = '说明'

    leave_data = [
        ('请假申请ID', 'CHAR(15)', 'PK', '格式: LV{学号后8位}{时间戳后5位}'),
        ('学号', 'CHAR(12)', 'FK', '关联students'),
        ('课程ID', 'CHAR(10)', 'FK', '关联courses'),
        ('请假类型', 'VARCHAR(20)', 'NOT NULL', '病假、事假、其他'),
        ('开始/结束日期', 'DATE', 'NOT NULL', '请假时间范围'),
        ('请假理由', 'TEXT', '', '详细理由'),
        ('病假条图片', 'VARCHAR(255)', '', '存储路径 (media/...)'),
        ('审批状态', 'TINYINT', 'NOT NULL', '0-待审批, 1-通过, 2-拒绝'),
        ('审批意见', 'VARCHAR(500)', '', '教师审批反馈'),
    ]
    for row in leave_data:
        cells = table2.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val

    # 4.2 关键数据表ER图
    add_heading_with_color(doc, '2. 关键数据表ER图', 2)
    doc.add_paragraph('考勤模块的实体关系设计如下：')
    doc.add_paragraph('• 教师 (teachers) <1:N> 课程 (courses)')
    doc.add_paragraph('• 学生 (students) <1:N> 请假申请 (leave_requests)')
    doc.add_paragraph('• 请假申请 (leave_requests) <1:N> 考勤记录 (attendances) [通过请假申请ID关联]')

    add_image_placeholder(doc, '考勤模块数据库ER图（重点展示attendances与leave_requests的关联）', RGBColor(255, 165, 0))

    # ================= 五、模块设计 =================
    add_heading_with_color(doc, '五、模块设计', 1)

    # 5.1 公共模块
    add_heading_with_color(doc, '1. 公共模块设计', 2)
    doc.add_paragraph('公共模块处理认证、工具函数及基础模板。')
    doc.add_paragraph('• 认证机制：使用 @require_login 装饰器和 Session 存储 (student_id/teacher_id)。')
    doc.add_paragraph('• 模板继承：定义 base.html, base_teacher.html, base_admin.html 实现UI统一。')

    # 5.2 学生模块
    add_heading_with_color(doc, '2. 学生模块设计', 2)
    doc.add_paragraph('学生模块主要文件：student_portal/views.py, templates/student/。')
    doc.add_paragraph('核心功能界面：')
    add_image_placeholder(doc, '学生端功能截图拼图（登录页、选课页、考勤查询页）', RGBColor(0, 100, 0))

    # 5.3 管理员模块
    add_heading_with_color(doc, '3. 后台管理员模块设计', 2)
    doc.add_paragraph('管理员模块主要文件：admin_portal/views.py。使用 ModelForm 处理复杂表单。')
    add_image_placeholder(doc, '管理员端功能截图拼图（用户管理、课程管理）', RGBColor(0, 0, 139))

    # 5.4 教师模块
    add_heading_with_color(doc, '4. 老师模块设计', 2)
    doc.add_paragraph('教师模块主要文件：teacher_portal/views.py。支持成绩Excel导入和考勤管理。')
    add_image_placeholder(doc, '教师端功能截图拼图（课程列表、成绩录入）', RGBColor(139, 0, 139))

    # 5.5 自主设计模块（考勤）
    add_heading_with_color(doc, '5. 设计一个模块，并进行开发（考勤管理模块）', 2)
    doc.add_paragraph('本节详细阐述自主设计的考勤管理模块的实现细节。该模块分为三个核心子功能。')

    # 5.5.1 教师点名
    add_heading_with_color(doc, '（1）教师点名功能 (Teacher Attendance)', 3)
    doc.add_paragraph('交互设计：采用创新的头像点击交互。点击学生头像可循环切换状态（出勤→迟到→早退→缺勤→请假）。')
    doc.add_paragraph('技术实现：')
    doc.add_paragraph('• 前端：JS维护状态数组，使用Fetch API发送批量JSON数据。')
    doc.add_paragraph(
        '• 后端：batch_attendance() 视图接收数据，生成唯一考勤ID (ATT+学号后6位+时间戳)，使用原生SQL批量插入。')
    add_image_placeholder(doc, '教师点名界面截图（展示头像列表和状态Tag）', RGBColor(255, 0, 0))
    add_image_placeholder(doc, '后端 batch_attendance 代码逻辑截图', RGBColor(139, 0, 0))

    # 5.5.2 学生请假
    add_heading_with_color(doc, '（2）学生请假申请功能 (Leave Request)', 3)
    doc.add_paragraph('交互设计：支持表单填写和图片预览。')
    doc.add_paragraph('技术实现：')
    doc.add_paragraph('• 前端：使用 FileReader API 实现图片上传前的本地预览。')
    doc.add_paragraph('• 后端：submit_leave_request() 视图处理 request.FILES，保存图片至 MEDIA_ROOT，生成请假ID (LV+...)。')
    add_image_placeholder(doc, '学生请假申请界面截图（含图片预览）', RGBColor(0, 0, 255))

    # 5.5.3 审批与自动化
    add_heading_with_color(doc, '（3）审批与考勤自动化 (Approval Logic)', 3)
    doc.add_paragraph('业务逻辑：审批不仅仅是状态更新，更触发数据联动。')
    doc.add_paragraph('核心算法：')
    doc.add_paragraph('1. 验证教师权限（是否教授该课程）。')
    doc.add_paragraph('2. 更新 leave_requests 表状态。')
    doc.add_paragraph('3. 若审批通过：遍历请假开始至结束日期的每一天，在 attendances 表中自动插入状态为“请假”(5) 的记录。')
    add_image_placeholder(doc, '请假审批模态框截图', RGBColor(0, 128, 0))
    add_image_placeholder(doc, '后端 update_leave_status 代码逻辑截图（展示自动创建考勤记录的部分）', RGBColor(0, 100, 0))

    # 保存文件
    file_name = '实验报告_第三部分_学生管理系统.docx'
    doc.save(file_name)
    print(f"成功生成文档：{file_name}")


if __name__ == '__main__':
    create_report()